# -*- coding: utf-8 -*-
"""
测试数据准备脚本
从多个文件提取问答对，构建评估测试集
"""

import json
import os
import random
from typing import List, Dict, Any
from docx import Document

# 配置
DATA_DIR = r"C:\Users\32459\Desktop\可用劳动数据集"
OUTPUT_FILE = r"D:\projects\fastapi_langchain_env\NAIVERAG_New\10.0_version\evaluate\test_dataset.json"

# 目标数量
TARGET_COUNT = {
    "jsonl": 100,      # 从JSONL提取
    "word_qa": 50,     # 从Word问答文件提取
    "word_law": 30,    # 从法规文件提取
    "word_dispute": 20 # 从纠纷文件提取
}


def load_jsonl_qa(file_path: str, max_count: int = 100) -> List[Dict]:
    """从JSONL文件提取问答对"""
    print(f"正在读取: {os.path.basename(file_path)}")
    
    qa_pairs = []
    with open(file_path, 'r', encoding='utf-8') as f:
        for i, line in enumerate(f):
            if i >= max_count:
                break
            try:
                data = json.loads(line.strip())
                qa_pairs.append({
                    "id": f"jsonl_{i}",
                    "query": data.get("input", ""),
                    "answer": data.get("output", ""),
                    "source": "jsonl",
                    "question_type": classify_question(data.get("input", ""))
                })
            except json.JSONDecodeError:
                continue
    
    print(f"  - 提取问答对: {len(qa_pairs)}条")
    return qa_pairs


def extract_qa_from_word(file_path: str, max_count: int = 50) -> List[Dict]:
    """从Word文档提取问答对"""
    print(f"正在读取: {os.path.basename(file_path)}")
    
    qa_pairs = []
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 尝试识别问答模式
        i = 0
        while i < len(paragraphs) and len(qa_pairs) < max_count:
            para = paragraphs[i]
            
            # 识别问题模式
            if is_question(para):
                # 查找对应的答案（通常是下一段或后续几段）
                answer = extract_answer(paragraphs, i)
                if answer:
                    qa_pairs.append({
                        "id": f"word_qa_{len(qa_pairs)}",
                        "query": clean_text(para),
                        "answer": answer,
                        "source": "word_qa",
                        "question_type": classify_question(para)
                    })
            i += 1
        
    except Exception as e:
        print(f"  - 读取失败: {e}")
    
    print(f"  - 提取问答对: {len(qa_pairs)}条")
    return qa_pairs


def extract_law_questions(file_path: str, max_count: int = 30) -> List[Dict]:
    """从法规文件提取问题"""
    print(f"正在读取: {os.path.basename(file_path)}")
    
    questions = []
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 基于法规内容生成问题
        for i, para in enumerate(paragraphs):
            if len(questions) >= max_count:
                break
            
            # 提取关键条款
            if any(keyword in para for keyword in ["第", "条", "规定", "应当", "必须", "不得"]):
                question = generate_question_from_law(para)
                if question:
                    questions.append({
                        "id": f"word_law_{len(questions)}",
                        "query": question,
                        "answer": para[:500],  # 使用原文作为答案参考
                        "source": "word_law",
                        "question_type": "complex"
                    })
                
    except Exception as e:
        print(f"  - 读取失败: {e}")
    
    print(f"  - 提取问题: {len(questions)}条")
    return questions


def extract_dispute_questions(file_path: str, max_count: int = 20) -> List[Dict]:
    """从纠纷文件提取问题"""
    print(f"正在读取: {os.path.basename(file_path)}")
    
    questions = []
    try:
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        for para in paragraphs:
            if len(questions) >= max_count:
                break
            
            # 识别问题模式
            if is_question(para) or any(kw in para for kw in ["怎么办", "如何", "怎么处理", "怎么处理"]):
                questions.append({
                    "id": f"word_dispute_{len(questions)}",
                    "query": clean_text(para),
                    "answer": "",
                    "source": "word_dispute",
                    "question_type": classify_question(para)
                })
                
    except Exception as e:
        print(f"  - 读取失败: {e}")
    
    print(f"  - 提取问题: {len(questions)}条")
    return questions


def is_question(text: str) -> bool:
    """判断文本是否为问题"""
    question_markers = ["？", "?", "吗", "怎么", "如何", "什么", "为什么", "请问", "能否", "是否可以"]
    return any(marker in text for marker in question_markers)


def extract_answer(paragraphs: List[str], question_idx: int) -> str:
    """提取问题的答案"""
    # 查找后续段落作为答案
    answer_parts = []
    for i in range(question_idx + 1, min(question_idx + 5, len(paragraphs))):
        para = paragraphs[i]
        # 如果遇到下一个问题，停止
        if is_question(para):
            break
        if len(para) > 10:  # 忽略过短的段落
            answer_parts.append(para)
    
    return " ".join(answer_parts[:3])  # 最多取3段


def generate_question_from_law(article_text: str) -> str:
    """从法规条款生成问题"""
    # 简单的规则生成
    if "应当" in article_text:
        subject = article_text.split("应当")[0][-20:] if "应当" in article_text else ""
        return f"{subject}应当怎样？"
    elif "不得" in article_text:
        subject = article_text.split("不得")[0][-20:] if "不得" in article_text else ""
        return f"什么情况下{subject}不得做？"
    elif "必须" in article_text:
        subject = article_text.split("必须")[0][-20:] if "必须" in article_text else ""
        return f"{subject}必须做什么？"
    return ""


def classify_question(query: str) -> str:
    """分类问题类型"""
    query = query.lower()
    
    # 复杂问题特征
    complex_markers = ["和", "与", "以及", "或者", "还是", "哪些", "分别", "同时", "而且"]
    if any(marker in query for marker in complex_markers):
        return "complex"
    
    # 模糊问题特征
    vague_markers = ["怎么办", "怎么处理", "有没有", "能不能", "好不好"]
    if any(marker in query for marker in vague_markers):
        return "vague"
    
    return "simple"


def clean_text(text: str) -> str:
    """清理文本"""
    # 移除多余空白
    text = " ".join(text.split())
    # 移除特殊字符
    text = text.replace("\u3000", " ")
    return text.strip()


def main():
    """主函数"""
    print("="*60)
    print("测试数据准备")
    print("="*60)
    
    all_test_cases = []
    
    # 1. 从JSONL文件提取问答对
    print("\n[1/4] 从JSONL文件提取问答对...")
    jsonl_file = os.path.join(DATA_DIR, "DISC-Law-SFT-Pair-QA-released.jsonl")
    if os.path.exists(jsonl_file):
        jsonl_data = load_jsonl_qa(jsonl_file, TARGET_COUNT["jsonl"])
        all_test_cases.extend(jsonl_data)
    
    # 2. 从Word问答文件提取
    print("\n[2/4] 从Word问答文件提取...")
    qa_files = [
        "劳动合同法常见问题解答汇编（200个）.docx",
    ]
    for filename in qa_files:
        file_path = os.path.join(DATA_DIR, filename)
        if os.path.exists(file_path):
            qa_data = extract_qa_from_word(file_path, TARGET_COUNT["word_qa"])
            all_test_cases.extend(qa_data)
    
    # 3. 从法规文件提取问题
    print("\n[3/4] 从法规文件提取问题...")
    law_files = [
        "职工带薪年休假条例实施细则.docx",
        "中华人民共和国职业病防治法.docx",
    ]
    for filename in law_files:
        file_path = os.path.join(DATA_DIR, filename)
        if os.path.exists(file_path):
            law_data = extract_law_questions(file_path, TARGET_COUNT["word_law"])
            all_test_cases.extend(law_data)
    
    # 4. 从纠纷文件提取问题
    print("\n[4/4] 从纠纷文件提取问题...")
    dispute_files = [
        "可能遇到的劳动纠纷问题.docx",
    ]
    for filename in dispute_files:
        file_path = os.path.join(DATA_DIR, filename)
        if os.path.exists(file_path):
            dispute_data = extract_dispute_questions(file_path, TARGET_COUNT["word_dispute"])
            all_test_cases.extend(dispute_data)
    
    # 随机打乱
    random.shuffle(all_test_cases)
    
    # 添加索引
    for i, case in enumerate(all_test_cases):
        case["id"] = f"test_{i:04d}"
    
    # 保存结果
    output_data = {
        "version": "1.0",
        "created_date": "2026-03-16",
        "total_count": len(all_test_cases),
        "test_cases": all_test_cases
    }
    
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    
    print("\n" + "="*60)
    print(f"测试数据准备完成!")
    print(f"  - 总数: {len(all_test_cases)}条")
    print(f"  - 输出文件: {OUTPUT_FILE}")
    print("="*60)
    
    # 统计
    source_stats = {}
    type_stats = {}
    for case in all_test_cases:
        source_stats[case["source"]] = source_stats.get(case["source"], 0) + 1
        type_stats[case["question_type"]] = type_stats.get(case["question_type"], 0) + 1
    
    print("\n数据统计:")
    print("  按来源:")
    for source, count in source_stats.items():
        print(f"    - {source}: {count}条")
    print("  按类型:")
    for qtype, count in type_stats.items():
        print(f"    - {qtype}: {count}条")


if __name__ == "__main__":
    main()
