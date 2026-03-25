"""
文档处理模块：支持PDF、TXT、MD格式的文档加载和处理
PDF提取采用混合引擎 + 逐级降级 + 表格提取 + OCR兜底 + 深度清洗
文本分割采用递归字符语义分割
"""
import time
import re
import os
from pathlib import Path
from typing import List, Dict

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# ---------- 尝试导入OCR相关库（增强版异常捕获）----------
try:
    from paddleocr import PaddleOCR
    import fitz  # PyMuPDF
    # 强制使用CPU模式，避免GPU显存/驱动问题
    ocr_reader = PaddleOCR(
        use_angle_cls=True,
        lang="ch",
        use_gpu=False,          # 稳定优先
        show_log=False,
        enable_mkldnn=True     # CPU加速
    )
    PADDLE_OCR_AVAILABLE = True
    print("✅ PaddleOCR 已加载 (CPU模式)")
except Exception as e:
    PADDLE_OCR_AVAILABLE = False
    print(f"⚠️ PaddleOCR 加载失败: {type(e).__name__}: {e}")
    print("   如需OCR功能，请确保安装正确：pip install paddlepaddle paddleocr")
    print("   或使用CPU模式：use_gpu=False")


class DocumentProcessor:
    """文档处理器，支持PDF、TXT、MD格式，PDF采用混合引擎降级提取+表格+OCR"""

    # ---------- 混合引擎优先级 ----------
    ENGINE_PRIORITY = [
        ("PyMuPDF", "_extract_with_pymupdf"),
        ("pdfplumber", "_extract_with_pdfplumber"),
        ("pypdf", "_extract_with_pypdf"),
        ("pdfminer", "_extract_with_pdfminer"),
    ]

    def __init__(self, chunk_size=300, chunk_overlap=50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # ========== 语义分割（强化） ==========
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",        # 多空行分割的大段落
                "\n\n",          # 段落
                "（", "）",      # 列表项（括号包围）
                "\n",            # 换行
                "。", "！", "？", # 句子
                "；",            # 分句
                "，",            # 短语
                " ",            # 单词
                ""              # 字符级回退
            ],
            keep_separator=False,   # 不保留分隔符，减少碎片
        )

    # ---------- 深度清洗函数（温和版）----------
    def _is_garbled(self, text: str, threshold: float = 0.3) -> bool:
        """检测文本是否乱码"""
        if not text or len(text) < 10:
            return False
        
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        readable_count = chinese_chars + len(re.findall(r'[a-zA-Z0-9\s\u3000-\u303f\uff00-\uffef]', text))
        
        if len(text) > 0:
            ratio = readable_count / len(text)
            return ratio < (1 - threshold)
        
        return False

    def _convert_pdf_via_images(self, input_path: str, output_path: str = None) -> str:
        """将PDF转为图片再合成新PDF（消除字体编码问题）"""
        import tempfile
        from PIL import Image
        
        if output_path is None:
            temp_file = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            output_path = temp_file.name
            temp_file.close()
        
        images = []
        with fitz.open(input_path) as doc:
            for page_num in range(len(doc)):
                page = doc[page_num]
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
        
        if images:
            images[0].save(
                output_path,
                save_all=True,
                append_images=images[1:],
                quality=95
            )
            print(f"   🔄 PDF已转换为图片格式")
        
        return output_path

    def _clean_text(self, text: str) -> str:
        """温和清除PDF提取中的控制字符和重复空格，保留有效内容"""
        if not text:
            return ""

        # 1. 移除控制字符和不可见字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        # 2. 修复常见PDF乱码（零宽空格、BOM等）
        text = text.replace('\ufeff', '')
        text = text.replace('\u200b', '')
        text = re.sub(r'[─━═■●◆]', '', text)       # 装饰性符号（可保留则注释）

        # 3. 合并多余空格和换行
        text = re.sub(r'[ \t]+', ' ', text)        # 多个空格→单个空格
        text = re.sub(r'\n\s*\n', '\n\n', text)    # 多个空行→最多两个换行

        # 4. 规范化中文引号（可选）
        text = text.replace('“', '"').replace('”', '"')
        text = text.replace('‘', "'").replace('’', "'")
        text = text.replace('（', '(').replace('）', ')')

        # 5. 去除每行首尾空格
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        # 6. 再次合并多余空格
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    # ---------- 各引擎提取函数（增加错误打印）----------
    def _extract_with_pymupdf(self, file_path: str) -> Dict[int, str]:
        try:
            import fitz
        except ImportError:
            return {}
        try:
            pages = {}
            with fitz.open(file_path) as doc:
                for page_num in range(len(doc)):
                    text = doc[page_num].get_text()
                    if text and text.strip():
                        cleaned = self._clean_text(text)
                        if cleaned:
                            pages[page_num + 1] = cleaned
            return pages
        except Exception as e:
            print(f"   ❌ PyMuPDF提取失败: {e}")
            return {}

    def _extract_with_pdfplumber(self, file_path: str) -> Dict[int, str]:
        try:
            import pdfplumber
        except ImportError:
            return {}
        try:
            pages = {}
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    native_text = page.extract_text() or ""
                    if native_text:
                        native_text = self._clean_text(native_text)

                    table_text = ""
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if any(cell for row in table for cell in row):
                                table_text += self._extract_table_to_text(table) + "\n"

                    combined = native_text + "\n" + table_text
                    combined = self._clean_text(combined)
                    if combined:
                        pages[page_num] = combined
            return pages
        except Exception as e:
            print(f"   ❌ pdfplumber提取失败: {e}")
            return {}

    def _extract_with_pypdf(self, file_path: str) -> Dict[int, str]:
        try:
            from pypdf import PdfReader
        except ImportError:
            return {}
        try:
            pages = {}
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    cleaned = self._clean_text(text)
                    if cleaned:
                        pages[page_num] = cleaned
            return pages
        except Exception as e:
            print(f"   ❌ pypdf提取失败: {e}")
            return {}

    def _extract_with_pdfminer(self, file_path: str) -> Dict[int, str]:
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            from io import StringIO
        except ImportError:
            return {}
        try:
            output = StringIO()
            with open(file_path, 'rb') as f:
                extract_text_to_fp(f, output, laparams=LAParams(), output_type='text', codec='utf-8')
            full_text = output.getvalue()
            if full_text and full_text.strip():
                cleaned = self._clean_text(full_text)
                if cleaned:
                    return {1: cleaned}
            return {}
        except Exception as e:
            print(f"   ❌ pdfminer提取失败: {e}")
            return {}

    # ---------- 表格转文本 ----------
    def _extract_table_to_text(self, table_data) -> str:
        """将pdfplumber表格二维列表转为Markdown风格文本"""
        if not table_data:
            return ""
        lines = []
        lines.append("【表格】")
        header = [str(cell).strip() if cell else "" for cell in table_data[0]]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table_data[1:]:
            row_cells = [str(cell).strip() if cell else "" for cell in row]
            lines.append("| " + " | ".join(row_cells) + " |")
        return "\n".join(lines)

    # ---------- OCR兜底（已稳定）----------
    def _ocr_page(self, pdf_path: str, page_num: int, retry: int = 2) -> tuple:
        """
        使用PaddleOCR识别整页PDF（CPU模式，稳定可靠）
        
        Returns:
            tuple: (提取的文本, 是否成功处理, 错误信息)
        """
        if not PADDLE_OCR_AVAILABLE:
            return "", False, "PaddleOCR未安装"
        
        last_error = None
        for attempt in range(1, retry + 1):
            try:
                doc = fitz.open(pdf_path)
                if page_num < 1 or page_num > len(doc):
                    doc.close()
                    return "", False, f"页码无效: {page_num}"
                
                page = doc[page_num - 1]
                pix = page.get_pixmap(dpi=300, colorspace="gray")
                
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    temp_img = tmp.name
                pix.save(temp_img)
                doc.close()

                result = ocr_reader.ocr(temp_img, cls=True)
                
                if os.path.exists(temp_img):
                    os.remove(temp_img)

                if result and len(result) > 0 and result[0] is not None:
                    ocr_text = "\n".join([line[1][0] for line in result[0]])
                    return self._clean_text(ocr_text), True, ""
                else:
                    return "", True, "OCR无结果"  # 处理正常但无内容
                    
            except Exception as e:
                last_error = e
                if attempt < retry:
                    time.sleep(0.5)  # 重试前等待
        
        error_msg = f"OCR第{page_num}页失败 (重试{retry}次): {last_error}"
        print(f"   ❌ {error_msg}")
        return "", False, str(last_error)

    # ---------- PDF加载主逻辑（独立获取页数 + 引擎缓存 + OCR）----------
    def load_pdf_document(self, file_path: str) -> List[Document]:
        """支持扫描件OCR的PDF加载方法"""
        documents = []
        print(f"📖 正在读取PDF文件: {Path(file_path).name} (混合引擎+表格+OCR)")

        # 1. 独立获取总页数（不依赖引擎提取结果）
        total_pages = 0
        try:
            import fitz
            with fitz.open(file_path) as doc:
                total_pages = len(doc)
            print(f"   ✅ PyMuPDF获取页数: {total_pages}")
        except:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                print(f"   ✅ pypdf获取页数: {total_pages}")
            except Exception as e:
                print(f"   ❌ 无法获取PDF页数: {e}")
                return []

        if total_pages == 0:
            print("   ⚠️ PDF页数为0，文件可能为空")
            return []

        # 2. 预先缓存各引擎的提取结果（提高性能）
        engine_pages_cache = {}
        for name, method_name in self.ENGINE_PRIORITY:
            try:
                method = getattr(self, method_name)
                pages = method(file_path)
                if pages:
                    engine_pages_cache[name] = pages
                    print(f"   ✅ 引擎 {name} 成功提取 {len(pages)} 页")
            except Exception as e:
                print(f"   ⚠️ 引擎 {name} 出错: {e}")

        # 3. 逐页处理（优先引擎 → 降级OCR → 乱码转换）
        failed_pages = []
        ocr_failed_pages = []
        
        for page_num in range(1, total_pages + 1):
            page_text = ""
            used_engine = None

            # 按优先级从缓存中查找该页文本
            for name, _ in self.ENGINE_PRIORITY:
                if name in engine_pages_cache and page_num in engine_pages_cache[name]:
                    page_text = engine_pages_cache[name][page_num]
                    used_engine = name
                    
                    # 检测乱码，如果乱码直接使用OCR
                    if self._is_garbled(page_text):
                        print(f"   ⚠️ 第{page_num}页检测到乱码，尝试OCR...")
                        if PADDLE_OCR_AVAILABLE:
                            ocr_text, ocr_success, ocr_error = self._ocr_page(file_path, page_num)
                            if ocr_text:
                                page_text = ocr_text
                                used_engine = "PaddleOCR"
                                print(f"   ✅ 第{page_num}页OCR成功")
                            elif ocr_error:
                                print(f"   ⚠️ 第{page_num}页 OCR 处理失败: {ocr_error}")
                            else:
                                print(f"   ⚠️ 第{page_num}页 OCR 无结果")
                        else:
                            print(f"   ⚠️ PaddleOCR未安装")
                    break

            # 如果所有引擎都失败，尝试OCR
            if not page_text:
                if PADDLE_OCR_AVAILABLE:
                    print(f"   🔍 第{page_num}页无原生文本，尝试OCR...")
                    ocr_text, ocr_success, ocr_error = self._ocr_page(file_path, page_num)
                    if ocr_text:
                        page_text = ocr_text
                        used_engine = "PaddleOCR"
                        print(f"   ✅ 第{page_num}页OCR成功")
                    elif ocr_error:
                        print(f"   ⚠️ 第{page_num}页 OCR 处理失败: {ocr_error}")
                    else:
                        print(f"   ⚠️ 第{page_num}页 OCR 无结果")
                else:
                    print(f"   ⚠️ 第{page_num}页无文本，且OCR未安装")

            if page_text:
                doc = Document(
                    page_content=page_text,
                    metadata={
                        "source": file_path,
                        "type": "pdf",
                        "page": page_num,
                        "total_pages": total_pages,
                        "extract_engine": used_engine
                    }
                )
                documents.append(doc)
                print(f"   ✅ 第{page_num}/{total_pages}页 (引擎: {used_engine})")
            else:
                failed_pages.append(page_num)
                print(f"   ❌ 第{page_num}/{total_pages}页 所有方法均失败")

        # 输出处理统计
        if not documents:
            print(f"⚠️  PDF文件 {Path(file_path).name} 未提取到任何文本")
        else:
            print(f"📄  PDF文件共提取 {len(documents)}/{total_pages} 页")
            if failed_pages:
                print(f"   ⚠️ 失败页面: {failed_pages}")
        
        return documents

    # ---------- TXT/MD处理（多编码自动检测）----------
    def load_txt_document(self, file_path: str) -> List[Document]:
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                if not self._is_garbled(text):
                    text = self._clean_text(text)
                    doc = Document(
                        page_content=text,
                        metadata={"source": file_path, "type": "text", "encoding": encoding}
                    )
                    return [doc]
            except UnicodeDecodeError:
                continue
        
        print(f"❌ 读取TXT文件失败（尝试了多种编码）: {file_path}")
        return []

    def load_md_document(self, file_path: str) -> List[Document]:
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                if not self._is_garbled(text):
                    text = self._clean_text(text)
                    doc = Document(
                        page_content=text,
                        metadata={"source": file_path, "type": "markdown", "encoding": encoding}
                    )
                    return [doc]
            except UnicodeDecodeError:
                continue
        
        print(f"❌ 读取MD文件失败（尝试了多种编码）: {file_path}")
        return []

    # ---------- Word文档处理 ----------
    def load_docx_document(self, file_path: str) -> List[Document]:
        """加载Word文档(.docx)"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            full_text = self._clean_text(full_text)
            
            doc = Document(
                page_content=full_text,
                metadata={"source": file_path, "type": "docx"}
            )
            print(f"   ✅ Word文档读取成功")
            return [doc]
        except Exception as e:
            print(f"❌ 读取Word文档失败 {file_path}: {e}")
            return []

    # ---------- Excel文档处理 ----------
    def load_excel_document(self, file_path: str) -> List[Document]:
        """加载Excel文档(.xlsx/.xls)"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, data_only=True)
            all_text = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                all_text.append(f"=== 工作表: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        all_text.append(" | ".join(row_text))
                
                all_text.append("")
            
            full_text = "\n".join(all_text)
            full_text = self._clean_text(full_text)
            
            doc = Document(
                page_content=full_text,
                metadata={"source": file_path, "type": "excel"}
            )
            print(f"   ✅ Excel文档读取成功")
            return [doc]
        except Exception as e:
            print(f"❌ 读取Excel文档失败 {file_path}: {e}")
            return []

    def load_documents(self, file_path: str) -> List[Document]:
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.pdf':
            return self.load_pdf_document(file_path)
        elif file_ext == '.txt':
            return self.load_txt_document(file_path)
        elif file_ext in ['.md', '.markdown']:
            return self.load_md_document(file_path)
        elif file_ext in ['.docx']:
            return self.load_docx_document(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.load_excel_document(file_path)
        else:
            print(f"⚠️  不支持的文件格式: {file_ext}，尝试按文本读取")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                text = self._clean_text(text)
                doc = Document(
                    page_content=text,
                    metadata={"source": file_path, "type": "unknown"}
                )
                return [doc]
            except Exception as e:
                print(f"❌ 无法读取文件: {file_path}, 错误: {e}")
                return []

    def process_file(self, file_path: str) -> List[Document]:
        all_chunks = []
        if not Path(file_path).exists():
            print(f"❌ 文件不存在: {file_path}")
            return all_chunks

        try:
            print(f"📄 处理文件: {Path(file_path).name}")
            docs = self.load_documents(file_path)
            if docs:
                chunks = self.text_splitter.split_documents(docs)
                for i, chunk in enumerate(chunks):
                    chunk.metadata.update({
                        "chunk_id": i,
                        "total_chunks": len(chunks),
                        "file_name": Path(file_path).name,
                        "file_type": Path(file_path).suffix.lower()[1:],
                        "processed_time": time.strftime("%Y-%m-%d %H:%M:%S")
                    })
                all_chunks.extend(chunks)
                print(f"✅ 文件处理完成，分成 {len(chunks)} 个语义块")
            else:
                print(f"❌ 文件为空或无法读取")
        except Exception as e:
            print(f"❌ 处理失败 {Path(file_path).name}: {e}")

        return all_chunks